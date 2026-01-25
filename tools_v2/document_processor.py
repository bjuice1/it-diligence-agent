"""
Document Processor - Orchestrates document processing pipeline

Connects uploaded documents to the analysis engine:
- Manages processing queue with priority
- Extracts content from various document types
- Tracks processing status and progress
- Integrates with fact extraction and merging

Steps 1-15 of Phase 1
"""

import logging
import hashlib
import os
import threading
import time
from dataclasses import dataclass, field
from datetime import datetime
from typing import Dict, List, Optional, Any, Callable
from enum import Enum
from queue import PriorityQueue
import json

logger = logging.getLogger(__name__)


class ProcessingStatus(Enum):
    """Status of document processing."""
    QUEUED = "queued"
    EXTRACTING = "extracting"
    ANALYZING = "analyzing"
    CLASSIFYING = "classifying"
    MERGING = "merging"
    COMPLETE = "complete"
    FAILED = "failed"


class ProcessingPriority(Enum):
    """Priority levels for processing queue."""
    URGENT = 1
    HIGH = 2
    NORMAL = 3
    LOW = 4


@dataclass
class ProcessingProgress:
    """Tracks progress of a document through the pipeline."""
    doc_id: str
    filename: str
    status: ProcessingStatus = ProcessingStatus.QUEUED
    priority: ProcessingPriority = ProcessingPriority.NORMAL

    # Progress tracking
    current_stage: str = "queued"
    progress_percent: int = 0
    stages_completed: List[str] = field(default_factory=list)

    # Timing
    queued_at: str = ""
    started_at: str = ""
    completed_at: str = ""

    # Results
    facts_extracted: int = 0
    facts_added: int = 0
    facts_updated: int = 0
    conflicts_found: int = 0

    # Error handling
    error_message: str = ""
    retry_count: int = 0
    max_retries: int = 3

    # Processing log
    log_entries: List[Dict[str, Any]] = field(default_factory=list)

    def add_log(self, message: str, level: str = "info"):
        """Add entry to processing log."""
        self.log_entries.append({
            "timestamp": datetime.now().isoformat(),
            "level": level,
            "message": message
        })

    def to_dict(self) -> Dict[str, Any]:
        return {
            "doc_id": self.doc_id,
            "filename": self.filename,
            "status": self.status.value,
            "priority": self.priority.value,
            "current_stage": self.current_stage,
            "progress_percent": self.progress_percent,
            "stages_completed": self.stages_completed,
            "queued_at": self.queued_at,
            "started_at": self.started_at,
            "completed_at": self.completed_at,
            "facts_extracted": self.facts_extracted,
            "facts_added": self.facts_added,
            "facts_updated": self.facts_updated,
            "conflicts_found": self.conflicts_found,
            "error_message": self.error_message,
            "retry_count": self.retry_count,
            "log_entries": self.log_entries
        }


@dataclass
class ExtractionResult:
    """Result of content extraction from a document."""
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    page_count: int = 0
    word_count: int = 0
    chunks: List[str] = field(default_factory=list)  # For large documents
    extraction_method: str = ""
    success: bool = True
    error: str = ""


class ContentExtractor:
    """
    Extracts text content from various document types.

    Supports: PDF, DOCX, XLSX, CSV, TXT, MD, JSON
    """

    # Maximum pages before chunking
    CHUNK_THRESHOLD = 50
    CHUNK_SIZE = 10  # pages per chunk

    def __init__(self):
        self.extractors = {
            'pdf': self._extract_pdf,
            'docx': self._extract_docx,
            'doc': self._extract_docx,  # Try docx parser
            'xlsx': self._extract_xlsx,
            'xls': self._extract_xlsx,
            'csv': self._extract_csv,
            'txt': self._extract_text,
            'text': self._extract_text,
            'md': self._extract_text,
            'markdown': self._extract_text,
            'json': self._extract_json,
        }

    def extract(self, file_path: str, doc_type: str = "") -> ExtractionResult:
        """
        Extract content from a document.

        Args:
            file_path: Path to the document
            doc_type: Document type (auto-detected if not provided)

        Returns:
            ExtractionResult with content and metadata
        """
        if not doc_type:
            doc_type = self._detect_type(file_path)

        extractor = self.extractors.get(doc_type)
        if not extractor:
            return ExtractionResult(
                content="",
                success=False,
                error=f"Unsupported document type: {doc_type}"
            )

        try:
            result = extractor(file_path)
            result.extraction_method = doc_type

            # Chunk large documents
            if result.page_count > self.CHUNK_THRESHOLD:
                result.chunks = self._chunk_content(result.content, result.page_count)

            return result

        except Exception as e:
            logger.error(f"Extraction failed for {file_path}: {e}")
            return ExtractionResult(
                content="",
                success=False,
                error=str(e)
            )

    def _detect_type(self, file_path: str) -> str:
        """Detect document type from file extension."""
        ext = os.path.splitext(file_path)[1].lower().lstrip('.')
        return ext if ext else 'unknown'

    def _extract_pdf(self, file_path: str) -> ExtractionResult:
        """Extract text from PDF using PyMuPDF."""
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(file_path)
            content_parts = []

            for page_num, page in enumerate(doc):
                text = page.get_text()
                content_parts.append(f"[Page {page_num + 1}]\n{text}")

            content = "\n\n".join(content_parts)

            return ExtractionResult(
                content=content,
                page_count=len(doc),
                word_count=len(content.split()),
                metadata={
                    "title": doc.metadata.get("title", ""),
                    "author": doc.metadata.get("author", ""),
                    "created": doc.metadata.get("creationDate", "")
                }
            )
        except ImportError:
            # Fallback to pdfplumber if PyMuPDF not available
            return self._extract_pdf_fallback(file_path)

    def _extract_pdf_fallback(self, file_path: str) -> ExtractionResult:
        """Fallback PDF extraction using pdfplumber."""
        try:
            import pdfplumber

            content_parts = []
            page_count = 0

            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    content_parts.append(f"[Page {i + 1}]\n{text}")

            content = "\n\n".join(content_parts)

            return ExtractionResult(
                content=content,
                page_count=page_count,
                word_count=len(content.split())
            )
        except Exception as e:
            return ExtractionResult(content="", success=False, error=str(e))

    def _extract_docx(self, file_path: str) -> ExtractionResult:
        """Extract text from DOCX."""
        try:
            from docx import Document

            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs]
            content = "\n\n".join(paragraphs)

            # Also extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                content += "\n\n[Table]\n" + "\n".join(table_text)

            return ExtractionResult(
                content=content,
                page_count=max(1, len(content) // 3000),  # Estimate
                word_count=len(content.split()),
                metadata={
                    "core_properties": str(doc.core_properties.__dict__) if hasattr(doc, 'core_properties') else ""
                }
            )
        except ImportError:
            return ExtractionResult(
                content="",
                success=False,
                error="python-docx not installed"
            )
        except Exception as e:
            return ExtractionResult(content="", success=False, error=str(e))

    def _extract_xlsx(self, file_path: str) -> ExtractionResult:
        """Extract text from Excel files."""
        try:
            import pandas as pd

            # Read all sheets
            xlsx = pd.ExcelFile(file_path)
            content_parts = []

            for sheet_name in xlsx.sheet_names:
                df = pd.read_excel(xlsx, sheet_name=sheet_name)
                content_parts.append(f"[Sheet: {sheet_name}]")
                content_parts.append(df.to_string())

            content = "\n\n".join(content_parts)

            return ExtractionResult(
                content=content,
                page_count=len(xlsx.sheet_names),
                word_count=len(content.split()),
                metadata={"sheets": xlsx.sheet_names}
            )
        except ImportError:
            return ExtractionResult(
                content="",
                success=False,
                error="pandas not installed"
            )
        except Exception as e:
            return ExtractionResult(content="", success=False, error=str(e))

    def _extract_csv(self, file_path: str) -> ExtractionResult:
        """Extract text from CSV files."""
        try:
            import pandas as pd

            df = pd.read_csv(file_path)
            content = df.to_string()

            return ExtractionResult(
                content=content,
                page_count=1,
                word_count=len(content.split()),
                metadata={
                    "columns": list(df.columns),
                    "rows": len(df)
                }
            )
        except Exception as e:
            # Fallback to simple read
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                return ExtractionResult(
                    content=content,
                    page_count=1,
                    word_count=len(content.split())
                )
            except Exception as e2:
                return ExtractionResult(content="", success=False, error=str(e2))

    def _extract_text(self, file_path: str) -> ExtractionResult:
        """Extract text from plain text/markdown files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            return ExtractionResult(
                content=content,
                page_count=max(1, len(content) // 3000),
                word_count=len(content.split())
            )
        except Exception as e:
            return ExtractionResult(content="", success=False, error=str(e))

    def _extract_json(self, file_path: str) -> ExtractionResult:
        """Extract and format JSON content."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            content = json.dumps(data, indent=2)

            return ExtractionResult(
                content=content,
                page_count=1,
                word_count=len(content.split()),
                metadata={"type": "json", "keys": list(data.keys()) if isinstance(data, dict) else []}
            )
        except Exception as e:
            return ExtractionResult(content="", success=False, error=str(e))

    def _chunk_content(self, content: str, page_count: int) -> List[str]:
        """Split large content into manageable chunks."""
        chunks = []

        # Try to split by page markers
        if "[Page " in content:
            pages = content.split("[Page ")
            current_chunk = []
            current_pages = 0

            for page in pages:
                if not page.strip():
                    continue
                current_chunk.append("[Page " + page if not page.startswith("[Page") else page)
                current_pages += 1

                if current_pages >= self.CHUNK_SIZE:
                    chunks.append("\n\n".join(current_chunk))
                    current_chunk = []
                    current_pages = 0

            if current_chunk:
                chunks.append("\n\n".join(current_chunk))
        else:
            # Simple character-based chunking
            chunk_size = len(content) // (page_count // self.CHUNK_SIZE + 1)
            for i in range(0, len(content), chunk_size):
                chunks.append(content[i:i + chunk_size])

        return chunks


class ProcessingQueue:
    """
    Priority-based queue for document processing.

    Features:
    - Priority ordering (urgent first)
    - FIFO within same priority
    - Thread-safe operations
    """

    def __init__(self):
        self._queue = PriorityQueue()
        self._lock = threading.Lock()
        self._counter = 0  # For FIFO ordering within priority
        self._active: Dict[str, ProcessingProgress] = {}
        self._completed: Dict[str, ProcessingProgress] = {}
        self._failed: Dict[str, ProcessingProgress] = {}

    def add(self, progress: ProcessingProgress) -> None:
        """Add document to processing queue."""
        with self._lock:
            self._counter += 1
            # Priority tuple: (priority_value, counter, progress)
            # Lower priority value = higher priority
            self._queue.put((progress.priority.value, self._counter, progress))
            progress.queued_at = datetime.now().isoformat()
            progress.add_log(f"Added to queue with priority {progress.priority.name}")

    def get_next(self) -> Optional[ProcessingProgress]:
        """Get next document for processing."""
        try:
            with self._lock:
                if self._queue.empty():
                    return None
                _, _, progress = self._queue.get_nowait()
                self._active[progress.doc_id] = progress
                return progress
        except:
            return None

    def mark_active(self, doc_id: str) -> None:
        """Mark document as actively processing."""
        if doc_id in self._active:
            progress = self._active[doc_id]
            progress.status = ProcessingStatus.EXTRACTING
            progress.started_at = datetime.now().isoformat()
            progress.add_log("Processing started")

    def mark_complete(self, doc_id: str) -> None:
        """Mark document as complete."""
        if doc_id in self._active:
            progress = self._active.pop(doc_id)
            progress.status = ProcessingStatus.COMPLETE
            progress.completed_at = datetime.now().isoformat()
            progress.progress_percent = 100
            progress.add_log("Processing complete")
            self._completed[doc_id] = progress

    def mark_failed(self, doc_id: str, error: str) -> None:
        """Mark document as failed."""
        if doc_id in self._active:
            progress = self._active[doc_id]
            progress.retry_count += 1
            progress.error_message = error
            progress.add_log(f"Processing failed: {error}", "error")

            if progress.retry_count < progress.max_retries:
                # Requeue for retry
                progress.status = ProcessingStatus.QUEUED
                progress.add_log(f"Requeuing for retry ({progress.retry_count}/{progress.max_retries})")
                self._active.pop(doc_id)
                self.add(progress)
            else:
                # Max retries exceeded
                progress.status = ProcessingStatus.FAILED
                progress.add_log("Max retries exceeded, marking as failed", "error")
                self._active.pop(doc_id)
                self._failed[doc_id] = progress

    def update_progress(self, doc_id: str, stage: str, percent: int) -> None:
        """Update processing progress."""
        if doc_id in self._active:
            progress = self._active[doc_id]
            progress.current_stage = stage
            progress.progress_percent = percent
            if stage not in progress.stages_completed:
                progress.stages_completed.append(stage)
            progress.add_log(f"Stage: {stage} ({percent}%)")

    def get_status(self, doc_id: str) -> Optional[ProcessingProgress]:
        """Get status for a document."""
        if doc_id in self._active:
            return self._active[doc_id]
        if doc_id in self._completed:
            return self._completed[doc_id]
        if doc_id in self._failed:
            return self._failed[doc_id]
        return None

    def get_queue_stats(self) -> Dict[str, Any]:
        """Get queue statistics."""
        return {
            "queued": self._queue.qsize(),
            "active": len(self._active),
            "completed": len(self._completed),
            "failed": len(self._failed)
        }

    def get_all_status(self) -> List[Dict[str, Any]]:
        """Get status of all documents."""
        all_progress = []

        # Get queued items (need to peek without removing)
        with self._lock:
            queued_items = list(self._queue.queue)
            for _, _, progress in queued_items:
                all_progress.append(progress.to_dict())

        # Active, completed, failed
        for progress in self._active.values():
            all_progress.append(progress.to_dict())
        for progress in self._completed.values():
            all_progress.append(progress.to_dict())
        for progress in self._failed.values():
            all_progress.append(progress.to_dict())

        return all_progress


class DocumentProcessor:
    """
    Main orchestrator for document processing pipeline.

    Manages the flow: Upload -> Extract -> Analyze -> Classify -> Merge
    """

    # Processing stages and their progress percentages
    STAGES = {
        "queued": 0,
        "extracting": 10,
        "analyzing": 30,
        "classifying": 60,
        "merging": 80,
        "complete": 100
    }

    def __init__(self, document_registry=None, fact_store=None, fact_merger=None):
        self.queue = ProcessingQueue()
        self.extractor = ContentExtractor()
        self.document_registry = document_registry
        self.fact_store = fact_store
        self.fact_merger = fact_merger

        # Background worker
        self._worker_thread = None
        self._stop_worker = threading.Event()
        self._worker_running = False

        # Callbacks for status updates
        self._status_callbacks: List[Callable] = []

        # Store for pending changes by tier (for UI)
        self.pending_changes: Dict[str, List[Dict]] = {
            "tier1": [],
            "tier2": [],
            "tier3": []
        }

        # Processing metrics
        self.metrics = {
            "total_processed": 0,
            "total_failed": 0,
            "avg_processing_time_ms": 0,
            "by_type": {}
        }

    def queue_document(
        self,
        doc_id: str,
        filename: str,
        file_path: str,
        priority: ProcessingPriority = ProcessingPriority.NORMAL
    ) -> ProcessingProgress:
        """
        Add a document to the processing queue.

        Args:
            doc_id: Document ID from registry
            filename: Original filename
            file_path: Path to the file
            priority: Processing priority

        Returns:
            ProcessingProgress for tracking
        """
        progress = ProcessingProgress(
            doc_id=doc_id,
            filename=filename,
            priority=priority
        )
        progress.add_log(f"Document queued: {filename}")

        # Store file path in metadata for processing
        progress.log_entries.append({
            "timestamp": datetime.now().isoformat(),
            "level": "debug",
            "message": f"file_path:{file_path}"
        })

        self.queue.add(progress)
        self._notify_status_change(progress)

        return progress

    def process_single(self, doc_id: str, file_path: str) -> ProcessingProgress:
        """
        Process a single document immediately (synchronous).

        Args:
            doc_id: Document ID
            file_path: Path to document file

        Returns:
            ProcessingProgress with results
        """
        progress = self.queue.get_status(doc_id)
        if not progress:
            progress = ProcessingProgress(
                doc_id=doc_id,
                filename=os.path.basename(file_path)
            )

        start_time = time.time()

        try:
            # Stage 1: Extract content
            self._update_stage(progress, "extracting", ProcessingStatus.EXTRACTING)
            extraction = self.extractor.extract(file_path)

            if not extraction.success:
                raise Exception(f"Extraction failed: {extraction.error}")

            progress.add_log(f"Extracted {extraction.word_count} words from {extraction.page_count} pages")

            # Stage 2: Analyze content (extract facts)
            self._update_stage(progress, "analyzing", ProcessingStatus.ANALYZING)
            facts = self._analyze_content(extraction, progress.filename)
            progress.facts_extracted = len(facts)
            progress.add_log(f"Extracted {len(facts)} facts")

            # Stage 3: Classify facts into tiers
            self._update_stage(progress, "classifying", ProcessingStatus.CLASSIFYING)
            classified = self._classify_facts(facts)
            progress.add_log(f"Classified facts: {len(classified.get('tier1', []))} auto-apply, "
                           f"{len(classified.get('tier2', []))} batch, "
                           f"{len(classified.get('tier3', []))} individual")

            # Store classified changes for review UI
            for tier in ["tier1", "tier2", "tier3"]:
                self.pending_changes[tier].extend(classified.get(tier, []))

            # Persist pending changes to disk
            self.save_pending_changes()

            # Stage 4: Merge only auto-eligible tier 1 facts
            # Tier 2 and 3 wait for human review
            self._update_stage(progress, "merging", ProcessingStatus.MERGING)

            # Get auto-eligible facts from tier 1
            tier1_facts = classified.get("tier1", [])
            auto_eligible_facts = [
                item["fact"] for item in tier1_facts
                if item.get("classification", {}).get("auto_apply_eligible", False)
            ]

            # Only merge auto-eligible facts
            merge_result = {"added": [], "updated": [], "conflicts": []}
            if auto_eligible_facts:
                merge_result = self._merge_facts(auto_eligible_facts, progress.filename)

            progress.facts_added = len(merge_result.get('added', []))
            progress.facts_updated = len(merge_result.get('updated', []))
            progress.conflicts_found = len(classified.get('tier3', []))  # Tier 3 = conflicts

            # Count facts pending review
            pending_review = len(classified.get('tier2', [])) + len(classified.get('tier3', []))
            non_auto_tier1 = len(tier1_facts) - len(auto_eligible_facts)
            pending_review += non_auto_tier1

            progress.add_log(f"Auto-merged: {progress.facts_added} added, "
                           f"{progress.facts_updated} updated. "
                           f"Pending review: {pending_review} facts")

            # Complete
            self._update_stage(progress, "complete", ProcessingStatus.COMPLETE)
            progress.completed_at = datetime.now().isoformat()

            # Update metrics
            elapsed_ms = int((time.time() - start_time) * 1000)
            self._update_metrics(progress, elapsed_ms, success=True)

            # Update document registry
            if self.document_registry:
                fact_ids = merge_result.get('added', []) + merge_result.get('updated', [])
                self.document_registry.mark_processed(doc_id, fact_ids, elapsed_ms)

            return progress

        except Exception as e:
            logger.error(f"Processing failed for {doc_id}: {e}")
            progress.status = ProcessingStatus.FAILED
            progress.error_message = str(e)
            progress.add_log(f"Processing failed: {e}", "error")

            self._update_metrics(progress, 0, success=False)

            if self.document_registry:
                self.document_registry.mark_failed(doc_id, str(e))

            return progress

    def _update_stage(self, progress: ProcessingProgress, stage: str, status: ProcessingStatus):
        """Update processing stage."""
        progress.current_stage = stage
        progress.status = status
        progress.progress_percent = self.STAGES.get(stage, 0)
        if stage not in progress.stages_completed:
            progress.stages_completed.append(stage)
        self._notify_status_change(progress)

    def _analyze_content(self, extraction: ExtractionResult, filename: str) -> List[Dict]:
        """
        Analyze content to extract facts using the incremental extractor.
        """
        try:
            from tools_v2.incremental_extractor import IncrementalExtractor

            extractor = IncrementalExtractor(fact_store=self.fact_store)
            context = extractor.build_context()

            # Extract facts from content
            extracted_facts = extractor.extract_facts_from_content(
                content=extraction.content,
                source_document=filename,
                context=context
            )

            # Convert to dict format for downstream processing
            return [fact.to_dict() for fact in extracted_facts]

        except Exception as e:
            logger.error(f"Fact extraction failed: {e}")
            return []

    def _classify_facts(self, facts: List[Dict]) -> Dict[str, List]:
        """
        Classify facts into review tiers.
        """
        try:
            from tools_v2.tier_classifier import TierClassifier, Tier

            classifier = TierClassifier(
                fact_store=self.fact_store,
                fact_merger=self.fact_merger
            )

            classified, stats = classifier.classify_batch(facts)

            # Convert Tier enum keys to strings
            return {
                "tier1": classified.get(Tier.AUTO_APPLY, []),
                "tier2": classified.get(Tier.BATCH_REVIEW, []),
                "tier3": classified.get(Tier.INDIVIDUAL_REVIEW, []),
                "stats": stats.to_dict()
            }

        except Exception as e:
            logger.error(f"Fact classification failed: {e}")
            return {
                "tier1": [],
                "tier2": [],
                "tier3": [],
                "stats": {}
            }

    def _merge_facts(self, facts: List[Dict], source_document: str) -> Dict[str, List]:
        """
        Merge facts into the fact store.

        Uses the FactMerger if available, otherwise placeholder.
        """
        if self.fact_merger:
            result = self.fact_merger.merge_document_facts(facts, source_document)
            return {
                "added": result.facts_added,
                "updated": result.facts_updated,
                "conflicts": [c.to_dict() for c in result.conflicts]
            }

        return {"added": [], "updated": [], "conflicts": []}

    def _update_metrics(self, progress: ProcessingProgress, elapsed_ms: int, success: bool):
        """Update processing metrics."""
        if success:
            self.metrics["total_processed"] += 1
            # Update rolling average
            total = self.metrics["total_processed"]
            current_avg = self.metrics["avg_processing_time_ms"]
            self.metrics["avg_processing_time_ms"] = (
                (current_avg * (total - 1) + elapsed_ms) / total
            )
        else:
            self.metrics["total_failed"] += 1

        # Track by type
        doc_type = os.path.splitext(progress.filename)[1].lower()
        if doc_type not in self.metrics["by_type"]:
            self.metrics["by_type"][doc_type] = {"processed": 0, "failed": 0}

        if success:
            self.metrics["by_type"][doc_type]["processed"] += 1
        else:
            self.metrics["by_type"][doc_type]["failed"] += 1

    # Background worker methods
    def start_worker(self):
        """Start background processing worker."""
        if self._worker_running:
            return

        self._stop_worker.clear()
        self._worker_thread = threading.Thread(target=self._worker_loop, daemon=True)
        self._worker_thread.start()
        self._worker_running = True
        logger.info("Document processing worker started")

    def stop_worker(self, wait: bool = True):
        """Stop background processing worker."""
        if not self._worker_running:
            return

        self._stop_worker.set()
        if wait and self._worker_thread:
            self._worker_thread.join(timeout=30)
        self._worker_running = False
        logger.info("Document processing worker stopped")

    def _worker_loop(self):
        """Background worker loop."""
        while not self._stop_worker.is_set():
            progress = self.queue.get_next()

            if progress:
                # Get file path from log entries
                file_path = None
                for entry in progress.log_entries:
                    if entry.get("message", "").startswith("file_path:"):
                        file_path = entry["message"].replace("file_path:", "")
                        break

                if file_path:
                    self.queue.mark_active(progress.doc_id)
                    result = self.process_single(progress.doc_id, file_path)

                    if result.status == ProcessingStatus.COMPLETE:
                        self.queue.mark_complete(progress.doc_id)
                    elif result.status == ProcessingStatus.FAILED:
                        self.queue.mark_failed(progress.doc_id, result.error_message)
            else:
                # No documents to process, wait a bit
                time.sleep(1)

    # Status callbacks
    def add_status_callback(self, callback: Callable):
        """Add callback for status updates."""
        self._status_callbacks.append(callback)

    def _notify_status_change(self, progress: ProcessingProgress):
        """Notify all callbacks of status change."""
        for callback in self._status_callbacks:
            try:
                callback(progress)
            except Exception as e:
                logger.error(f"Status callback error: {e}")

    def get_status(self, doc_id: str) -> Optional[Dict[str, Any]]:
        """Get processing status for a document."""
        progress = self.queue.get_status(doc_id)
        return progress.to_dict() if progress else None

    def get_all_status(self) -> List[Dict[str, Any]]:
        """Get status of all documents in queue."""
        return self.queue.get_all_status()

    def get_queue_stats(self) -> Dict[str, Any]:
        """Get queue statistics."""
        stats = self.queue.get_queue_stats()
        stats["metrics"] = self.metrics
        return stats

    # Pending changes persistence methods
    def save_pending_changes(self, output_dir: str = None) -> str:
        """
        Save pending changes to a file for persistence across session reloads.

        Args:
            output_dir: Directory to save to (defaults to data directory)

        Returns:
            Path to saved file
        """
        if output_dir is None:
            from config_v2 import OUTPUT_DIR
            output_dir = str(OUTPUT_DIR)

        filepath = os.path.join(output_dir, "pending_changes.json")

        try:
            with open(filepath, 'w') as f:
                json.dump(self.pending_changes, f, indent=2, default=str)
            logger.info(f"Saved pending changes to {filepath}")
            return filepath
        except Exception as e:
            logger.error(f"Failed to save pending changes: {e}")
            return ""

    def load_pending_changes(self, output_dir: str = None) -> bool:
        """
        Load pending changes from file.

        Args:
            output_dir: Directory to load from (defaults to data directory)

        Returns:
            True if loaded successfully
        """
        if output_dir is None:
            from config_v2 import OUTPUT_DIR
            output_dir = str(OUTPUT_DIR)

        filepath = os.path.join(output_dir, "pending_changes.json")

        if not os.path.exists(filepath):
            logger.debug(f"No pending changes file found at {filepath}")
            return False

        try:
            with open(filepath, 'r') as f:
                loaded = json.load(f)

            # Merge loaded changes with existing (avoid duplicates)
            for tier in ["tier1", "tier2", "tier3"]:
                existing_ids = {c.get("fact", {}).get("temp_id") for c in self.pending_changes.get(tier, [])}
                for change in loaded.get(tier, []):
                    temp_id = change.get("fact", {}).get("temp_id")
                    if temp_id and temp_id not in existing_ids:
                        self.pending_changes[tier].append(change)
                        existing_ids.add(temp_id)

            total = sum(len(self.pending_changes[t]) for t in ["tier1", "tier2", "tier3"])
            logger.info(f"Loaded {total} pending changes from {filepath}")
            return True
        except Exception as e:
            logger.error(f"Failed to load pending changes: {e}")
            return False

    def remove_pending_change(self, change_id: str) -> bool:
        """
        Remove a pending change by its temp_id and save to disk.

        Args:
            change_id: The temp_id of the change to remove

        Returns:
            True if removed successfully
        """
        removed = False
        for tier in ["tier1", "tier2", "tier3"]:
            original_len = len(self.pending_changes[tier])
            self.pending_changes[tier] = [
                c for c in self.pending_changes[tier]
                if c.get("fact", {}).get("temp_id") != change_id
            ]
            if len(self.pending_changes[tier]) < original_len:
                removed = True
                break

        if removed:
            self.save_pending_changes()

        return removed
