"""
Word Document Parser

Parses Word documents (.docx) extracting:
- Tables (converted to ParsedTable)
- Prose text (remaining content between/around tables)
- Document metadata
"""

import logging
import re
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass

from tools_v2.parsers.models import ParsedTable, ParseResult
from tools_v2.parsers.type_detector import detect_inventory_type

logger = logging.getLogger(__name__)

# Try to import python-docx for Word parsing
try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. Word parsing disabled. Install with: pip install python-docx")


def parse_word_document(file_path: Path) -> ParseResult:
    """
    Parse a Word document into tables and prose.

    Extracts all tables as ParsedTable objects and collects
    remaining text as prose content.

    Args:
        file_path: Path to .docx file

    Returns:
        ParseResult with tables, prose, and metadata
    """
    file_path = Path(file_path)
    result = ParseResult(
        source_file=file_path.name,
        file_type="word",
    )

    if not file_path.exists():
        result.errors.append(f"File not found: {file_path}")
        return result

    if not DOCX_AVAILABLE:
        result.errors.append(
            "python-docx not installed. Install with: pip install python-docx"
        )
        return result

    try:
        doc = Document(file_path)

        # Extract metadata
        result.metadata = _extract_metadata(doc)

        # Process document - tables and prose
        prose_parts = []
        table_index = 0

        # Iterate through document body in order
        for element in doc.element.body:
            # Check if it's a table
            if element.tag.endswith('tbl'):
                # Find the corresponding table object
                for table in doc.tables:
                    if table._tbl == element:
                        parsed_table = _parse_table(
                            table, file_path.name, table_index
                        )
                        if parsed_table and parsed_table.row_count > 0:
                            # Detect inventory type
                            inv_type, confidence = detect_inventory_type(parsed_table.headers)
                            parsed_table.detected_type = inv_type
                            parsed_table.detection_confidence = confidence

                            result.tables.append(parsed_table)

                            # Add placeholder in prose
                            prose_parts.append(f"\n[TABLE {table_index}: {inv_type}]\n")
                            table_index += 1
                        break

            # Check if it's a paragraph
            elif element.tag.endswith('p'):
                for para in doc.paragraphs:
                    if para._p == element:
                        text = para.text.strip()
                        if text:
                            prose_parts.append(text)
                        break

        # Combine prose
        result.prose = _clean_prose("\n".join(prose_parts))

    except Exception as e:
        result.errors.append(f"Failed to parse Word document: {str(e)}")
        logger.exception(f"Word parse error: {file_path}")

    return result


def _parse_table(table: 'Table', source_file: str, table_index: int) -> Optional[ParsedTable]:
    """Parse a docx Table into ParsedTable."""
    try:
        rows_data = []

        # Extract all rows
        for row in table.rows:
            cells = [_extract_cell_text(cell) for cell in row.cells]
            rows_data.append(cells)

        if not rows_data:
            return None

        # First non-empty row is assumed to be headers
        header_row_idx = 0
        for idx, row in enumerate(rows_data):
            if any(cell.strip() for cell in row):
                header_row_idx = idx
                break

        # Extract headers
        raw_headers = rows_data[header_row_idx]
        headers = [_normalize_header(h, i) for i, h in enumerate(raw_headers)]

        # Check if headers look valid
        if not _looks_like_headers(headers, rows_data[header_row_idx]):
            # Maybe first row is a title, try second row
            if len(rows_data) > header_row_idx + 1:
                alt_headers = [_normalize_header(h, i) for i, h in enumerate(rows_data[header_row_idx + 1])]
                if _looks_like_headers(alt_headers, rows_data[header_row_idx + 1]):
                    headers = alt_headers
                    header_row_idx += 1

        # Extract data rows
        rows = []
        for row_data in rows_data[header_row_idx + 1:]:
            if len(row_data) != len(headers):
                # Handle merged cells by padding/truncating
                if len(row_data) < len(headers):
                    row_data = row_data + [''] * (len(headers) - len(row_data))
                else:
                    row_data = row_data[:len(headers)]

            # Create row dict
            row_dict = {}
            has_data = False
            for i, value in enumerate(row_data):
                cleaned = _clean_cell_value(value)
                row_dict[headers[i]] = cleaned
                if cleaned:
                    has_data = True

            # Skip empty rows
            if has_data:
                rows.append(row_dict)

        if not rows:
            return None

        return ParsedTable(
            headers=headers,
            rows=rows,
            source_file=source_file,
            table_index=table_index,
        )

    except Exception as e:
        logger.warning(f"Failed to parse table {table_index}: {e}")
        return None


def _extract_cell_text(cell) -> str:
    """Extract text from a table cell, handling nested content."""
    text_parts = []

    for paragraph in cell.paragraphs:
        text = paragraph.text.strip()
        if text:
            text_parts.append(text)

    return " ".join(text_parts)


def _normalize_header(value: str, col_num: int) -> str:
    """Normalize a header value to lowercase string."""
    if not value or not value.strip():
        return f"column_{col_num}"

    header = value.strip().lower()

    # Remove special characters but keep spaces and underscores
    header = ''.join(c if c.isalnum() or c.isspace() or c == '_' else ' ' for c in header)

    # Collapse multiple spaces
    header = ' '.join(header.split())

    # Remove common citation markers from ToltIQ
    header = re.sub(r'\[\d+\]', '', header)
    header = re.sub(r'[\ue200\ue201\ue202]', '', header)

    header = header.strip()

    if not header:
        return f"column_{col_num}"

    return header


def _looks_like_headers(headers: List[str], raw_values: List[str]) -> bool:
    """Check if a row looks like headers vs data."""
    if not headers:
        return False

    # Headers should be mostly text, not numbers
    text_count = 0
    for raw in raw_values:
        if raw and raw.strip():
            # Check if it's primarily text (not just a number)
            cleaned = raw.strip()
            try:
                # Try to parse as number
                cleaned_num = cleaned.replace(",", "").replace("$", "").replace("%", "")
                float(cleaned_num)
                # It's a number
            except (ValueError, AttributeError):
                # Not a number - count as text
                text_count += 1

    # If most values are text, likely headers
    non_empty = sum(1 for v in raw_values if v and v.strip())
    return non_empty > 0 and text_count >= non_empty * 0.5


def _clean_cell_value(value: str) -> Optional[str]:
    """Clean a cell value."""
    if not value:
        return None

    cleaned = value.strip()

    # Remove citation markers
    cleaned = re.sub(r'\[\d+\]', '', cleaned)
    cleaned = re.sub(r'[\ue200\ue201\ue202]', '', cleaned)
    cleaned = cleaned.strip()

    if not cleaned or cleaned.lower() in ['n/a', 'na', '-', '--', 'tbd', 'none']:
        return None

    return cleaned


def _clean_prose(text: str) -> str:
    """Clean prose text - remove excessive whitespace and artifacts."""
    # Remove citation markers
    text = re.sub(r'\[\d+\]', '', text)
    text = re.sub(r'[\ue200\ue201\ue202]', '', text)

    # Collapse multiple newlines
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Remove lines that are just table placeholders if they're redundant
    lines = []
    for line in text.split('\n'):
        line = line.strip()
        if line:
            lines.append(line)

    return '\n'.join(lines)


def _extract_metadata(doc: 'Document') -> Dict[str, Any]:
    """Extract document metadata."""
    metadata = {}

    try:
        core_props = doc.core_properties

        if core_props.title:
            metadata['title'] = core_props.title
        if core_props.author:
            metadata['author'] = core_props.author
        if core_props.created:
            metadata['created'] = str(core_props.created)
        if core_props.modified:
            metadata['modified'] = str(core_props.modified)
        if core_props.subject:
            metadata['subject'] = core_props.subject

    except Exception as e:
        logger.debug(f"Could not extract metadata: {e}")

    # Count paragraphs and tables
    metadata['paragraph_count'] = len(doc.paragraphs)
    metadata['table_count'] = len(doc.tables)

    return metadata


def parse_docx_tables_only(file_path: Path) -> List[ParsedTable]:
    """
    Parse only tables from a Word document (faster, no prose extraction).

    Args:
        file_path: Path to .docx file

    Returns:
        List of ParsedTable objects
    """
    result = parse_word_document(file_path)
    return result.tables
