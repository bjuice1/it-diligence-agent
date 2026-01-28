"""
Document Store for IT Due Diligence System

Central registry for all source documents. Provides:
- SHA-256 hash-based integrity verification
- Entity separation (target vs buyer)
- Authority level tracking
- Document-to-fact traceability

Key Principle: Documents are the source of truth. Facts are interpretations.
The document hash is the immutable anchor for audit trails.
"""

import json
import shutil
import logging
from dataclasses import dataclass, field, asdict
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Dict, List, Optional, Any, Union
from uuid import uuid4
import threading

from utils.file_hasher import compute_sha256, compute_sha256_from_bytes

logger = logging.getLogger(__name__)


class DocumentStatus(Enum):
    """Document processing status."""
    PENDING = "pending"           # Uploaded, not yet processed
    EXTRACTING = "extracting"     # Text extraction in progress
    PROCESSED = "processed"       # Ready for discovery
    ERROR = "error"               # Processing failed


class AuthorityLevel(Enum):
    """
    Document authority levels for evidence weighting.

    Level 1 (highest): Official data room documents
    Level 2 (medium): Formal correspondence, contracts
    Level 3 (lowest): Discussion notes, informal communications
    """
    DATA_ROOM = 1
    CORRESPONDENCE = 2
    DISCUSSION_NOTES = 3


class Entity(Enum):
    """Entity that the document describes."""
    TARGET = "target"   # Company being acquired
    BUYER = "buyer"     # Acquiring company


@dataclass
class Document:
    """
    Represents a source document in the system.

    The document record is immutable once created (except for status).
    The hash ensures we can always verify the original file hasn't changed.
    """
    doc_id: str                           # UUID - system generated, immutable
    filename: str                         # Original filename as uploaded
    hash_sha256: str                      # Cryptographic hash of raw file bytes
    entity: str                           # "target" or "buyer"
    authority_level: int                  # 1 (high) to 3 (low)
    ingestion_timestamp: str              # ISO format timestamp
    uploaded_by: str                      # User/system that uploaded
    raw_file_path: str                    # Path to original unchanged file
    extracted_text_path: str              # Path to extracted text version
    page_count: int                       # Number of pages (0 if unknown)
    status: str                           # DocumentStatus value

    # Optional metadata
    file_size_bytes: int = 0              # Size of original file
    mime_type: str = ""                   # Detected MIME type
    extraction_error: str = ""            # Error message if status is ERROR

    # Tracking
    facts_extracted: int = 0              # Count of facts from this document
    last_discovery_run: str = ""          # Timestamp of last discovery

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for JSON serialization."""
        return asdict(self)

    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> "Document":
        """Create Document from dictionary."""
        return cls(**data)


class DocumentStore:
    """
    Central storage and registry for source documents.

    Thread-safe singleton pattern (like FactStore).

    Usage:
        store = DocumentStore.get_instance(base_dir=Path("output/documents"))

        # Add document
        doc = store.add_document(
            file_path="/path/to/upload.pdf",
            entity="target",
            authority_level=1,
            uploaded_by="user@example.com"
        )

        # Check for duplicate
        existing = store.get_by_hash(some_hash)

        # Get all target documents
        target_docs = store.get_documents_for_entity("target")
    """

    _instance: Optional["DocumentStore"] = None
    _lock = threading.Lock()

    def __init__(self, base_dir: Optional[Path] = None):
        """
        Initialize DocumentStore.

        Args:
            base_dir: Base directory for document storage.
                      Defaults to OUTPUT_DIR/documents from config.
        """
        if base_dir is None:
            try:
                from config_v2 import DOCUMENTS_DIR
                base_dir = DOCUMENTS_DIR
            except ImportError:
                base_dir = Path("output/documents")

        self.base_dir = Path(base_dir)
        self._documents: Dict[str, Document] = {}  # doc_id -> Document
        self._hash_index: Dict[str, str] = {}      # hash -> doc_id (for duplicate detection)
        self._lock = threading.Lock()

        # Create directory structure
        self._ensure_directories()

        # Load existing manifest if present
        self._load_manifest()

        logger.info(f"DocumentStore initialized at {self.base_dir}")

    @classmethod
    def get_instance(cls, base_dir: Optional[Path] = None) -> "DocumentStore":
        """Get singleton instance of DocumentStore."""
        with cls._lock:
            if cls._instance is None:
                cls._instance = cls(base_dir)
            return cls._instance

    @classmethod
    def reset_instance(cls):
        """Reset singleton (for testing)."""
        with cls._lock:
            cls._instance = None

    def _ensure_directories(self):
        """Create the document storage directory structure."""
        # Main directories
        self.base_dir.mkdir(parents=True, exist_ok=True)

        # Entity directories
        for entity in Entity:
            entity_dir = self.base_dir / entity.value
            entity_dir.mkdir(exist_ok=True)

            # Authority level subdirectories
            (entity_dir / "data_room").mkdir(exist_ok=True)
            (entity_dir / "correspondence").mkdir(exist_ok=True)
            (entity_dir / "notes").mkdir(exist_ok=True)

            # Extracted text directory
            (entity_dir / "extracted").mkdir(exist_ok=True)

        # Manifest file location
        self.manifest_path = self.base_dir / "manifest.json"

        logger.debug(f"Document directories created at {self.base_dir}")

    def _get_authority_folder(self, authority_level: int) -> str:
        """Get folder name for authority level."""
        mapping = {
            1: "data_room",
            2: "correspondence",
            3: "notes"
        }
        return mapping.get(authority_level, "notes")

    def add_document(
        self,
        file_path: Union[str, Path],
        entity: str,
        authority_level: int = 1,
        uploaded_by: str = "system"
    ) -> Document:
        """
        Add a document to the store.

        Args:
            file_path: Path to the source file
            entity: "target" or "buyer"
            authority_level: 1 (data room), 2 (correspondence), 3 (notes)
            uploaded_by: User or system identifier

        Returns:
            Document record (new or existing if duplicate)

        Raises:
            ValueError: If entity is invalid
            FileNotFoundError: If file doesn't exist
        """
        file_path = Path(file_path)

        # Validate entity
        if entity not in ("target", "buyer"):
            raise ValueError(
                f"Invalid entity '{entity}'. Must be 'target' or 'buyer'. "
                "Entity is determined by upload location, not inference."
            )

        # Validate authority level
        if authority_level not in (1, 2, 3):
            logger.warning(f"Invalid authority level {authority_level}, defaulting to 3")
            authority_level = 3

        # Compute hash
        file_hash = compute_sha256(file_path)

        # Check for duplicate
        with self._lock:
            if file_hash in self._hash_index:
                existing_doc_id = self._hash_index[file_hash]
                existing_doc = self._documents[existing_doc_id]
                logger.info(
                    f"Duplicate document detected: {file_path.name} "
                    f"matches existing {existing_doc.filename} (doc_id: {existing_doc_id})"
                )
                return existing_doc

        # Generate new document ID
        doc_id = str(uuid4())

        # Determine storage paths
        authority_folder = self._get_authority_folder(authority_level)
        storage_dir = self.base_dir / entity / authority_folder

        # Create unique filename: {short_hash}_{original_name}
        short_hash = file_hash[:8]
        safe_filename = f"{short_hash}_{file_path.name}"
        raw_file_path = storage_dir / safe_filename

        # Copy file to storage (preserving original)
        shutil.copy2(file_path, raw_file_path)
        logger.debug(f"Copied {file_path.name} to {raw_file_path}")

        # Extracted text path (will be populated during extraction)
        extracted_dir = self.base_dir / entity / "extracted"
        extracted_text_path = extracted_dir / f"{doc_id}.txt"

        # Get file size
        file_size = file_path.stat().st_size

        # Detect MIME type
        mime_type = self._detect_mime_type(file_path)

        # Create document record
        doc = Document(
            doc_id=doc_id,
            filename=file_path.name,
            hash_sha256=file_hash,
            entity=entity,
            authority_level=authority_level,
            ingestion_timestamp=datetime.now().isoformat(),
            uploaded_by=uploaded_by,
            raw_file_path=str(raw_file_path),
            extracted_text_path=str(extracted_text_path),
            page_count=0,  # Set during extraction
            status=DocumentStatus.PENDING.value,
            file_size_bytes=file_size,
            mime_type=mime_type
        )

        # Store document
        with self._lock:
            self._documents[doc_id] = doc
            self._hash_index[file_hash] = doc_id

        # Save manifest
        self._save_manifest()

        logger.info(
            f"Added document: {doc.filename} (doc_id: {doc_id}, "
            f"entity: {entity}, authority: {authority_level})"
        )

        return doc

    def add_document_from_bytes(
        self,
        file_bytes: bytes,
        filename: str,
        entity: str,
        authority_level: int = 1,
        uploaded_by: str = "system"
    ) -> Document:
        """
        Add a document from bytes (for web uploads).

        Args:
            file_bytes: Raw file content
            filename: Original filename
            entity: "target" or "buyer"
            authority_level: 1-3
            uploaded_by: User identifier

        Returns:
            Document record
        """
        # Validate entity
        if entity not in ("target", "buyer"):
            raise ValueError(
                f"Invalid entity '{entity}'. Must be 'target' or 'buyer'."
            )

        # Compute hash from bytes
        file_hash = compute_sha256_from_bytes(file_bytes)

        # Check for duplicate
        with self._lock:
            if file_hash in self._hash_index:
                existing_doc_id = self._hash_index[file_hash]
                existing_doc = self._documents[existing_doc_id]
                logger.info(f"Duplicate document: {filename} matches {existing_doc.filename}")
                return existing_doc

        # Generate doc ID and paths
        doc_id = str(uuid4())
        authority_folder = self._get_authority_folder(authority_level)
        storage_dir = self.base_dir / entity / authority_folder

        short_hash = file_hash[:8]
        safe_filename = f"{short_hash}_{filename}"
        raw_file_path = storage_dir / safe_filename

        # Write bytes to file
        with open(raw_file_path, "wb") as f:
            f.write(file_bytes)

        # Extracted text path
        extracted_dir = self.base_dir / entity / "extracted"
        extracted_text_path = extracted_dir / f"{doc_id}.txt"

        # Detect MIME type from extension
        mime_type = self._detect_mime_type(Path(filename))

        # Create document record
        doc = Document(
            doc_id=doc_id,
            filename=filename,
            hash_sha256=file_hash,
            entity=entity,
            authority_level=authority_level,
            ingestion_timestamp=datetime.now().isoformat(),
            uploaded_by=uploaded_by,
            raw_file_path=str(raw_file_path),
            extracted_text_path=str(extracted_text_path),
            page_count=0,
            status=DocumentStatus.PENDING.value,
            file_size_bytes=len(file_bytes),
            mime_type=mime_type
        )

        # Store
        with self._lock:
            self._documents[doc_id] = doc
            self._hash_index[file_hash] = doc_id

        self._save_manifest()

        logger.info(f"Added document from bytes: {filename} (doc_id: {doc_id})")
        return doc

    def _detect_mime_type(self, file_path: Path) -> str:
        """Detect MIME type from file extension."""
        extension_map = {
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ".csv": "text/csv",
            ".txt": "text/plain",
            ".md": "text/markdown"
        }
        ext = file_path.suffix.lower()
        return extension_map.get(ext, "application/octet-stream")

    def get_document(self, doc_id: str) -> Optional[Document]:
        """Get document by ID."""
        with self._lock:
            return self._documents.get(doc_id)

    def get_by_hash(self, file_hash: str) -> Optional[Document]:
        """
        Get document by its hash (for duplicate detection).

        Args:
            file_hash: SHA-256 hash of file

        Returns:
            Document if found, None otherwise
        """
        with self._lock:
            doc_id = self._hash_index.get(file_hash)
            if doc_id:
                return self._documents.get(doc_id)
            return None

    def get_documents_for_entity(self, entity: str) -> List[Document]:
        """
        Get all documents for a specific entity.

        Args:
            entity: "target" or "buyer"

        Returns:
            List of documents for that entity
        """
        with self._lock:
            return [
                doc for doc in self._documents.values()
                if doc.entity == entity
            ]

    def get_high_authority_docs(self, entity: Optional[str] = None) -> List[Document]:
        """
        Get documents with high authority level (data room).

        Args:
            entity: Optional entity filter

        Returns:
            List of authority level 1 documents
        """
        with self._lock:
            docs = [
                doc for doc in self._documents.values()
                if doc.authority_level == 1
            ]
            if entity:
                docs = [d for d in docs if d.entity == entity]
            return docs

    def get_discussion_notes(self, entity: Optional[str] = None) -> List[Document]:
        """
        Get discussion notes (lowest authority).

        Args:
            entity: Optional entity filter

        Returns:
            List of authority level 3 documents
        """
        with self._lock:
            docs = [
                doc for doc in self._documents.values()
                if doc.authority_level == 3
            ]
            if entity:
                docs = [d for d in docs if d.entity == entity]
            return docs

    def list_documents(
        self,
        entity: Optional[str] = None,
        authority_level: Optional[int] = None,
        status: Optional[str] = None
    ) -> List[Document]:
        """
        List documents with optional filters.

        Args:
            entity: Filter by entity
            authority_level: Filter by authority level
            status: Filter by status

        Returns:
            List of matching documents
        """
        with self._lock:
            docs = list(self._documents.values())

        if entity:
            docs = [d for d in docs if d.entity == entity]
        if authority_level:
            docs = [d for d in docs if d.authority_level == authority_level]
        if status:
            docs = [d for d in docs if d.status == status]

        # Sort by ingestion timestamp (newest first)
        docs.sort(key=lambda d: d.ingestion_timestamp, reverse=True)

        return docs

    def update_status(
        self,
        doc_id: str,
        status: DocumentStatus,
        error_message: str = ""
    ):
        """
        Update document processing status.

        Args:
            doc_id: Document ID
            status: New status
            error_message: Error details if status is ERROR
        """
        with self._lock:
            if doc_id not in self._documents:
                raise ValueError(f"Document not found: {doc_id}")

            doc = self._documents[doc_id]
            doc.status = status.value
            if error_message:
                doc.extraction_error = error_message

        self._save_manifest()
        logger.debug(f"Updated document {doc_id} status to {status.value}")

    def update_extraction_result(
        self,
        doc_id: str,
        page_count: int,
        extracted_text_path: Optional[str] = None
    ):
        """
        Update document after text extraction.

        Args:
            doc_id: Document ID
            page_count: Number of pages extracted
            extracted_text_path: Path to extracted text file
        """
        with self._lock:
            if doc_id not in self._documents:
                raise ValueError(f"Document not found: {doc_id}")

            doc = self._documents[doc_id]
            doc.page_count = page_count
            doc.status = DocumentStatus.PROCESSED.value
            if extracted_text_path:
                doc.extracted_text_path = extracted_text_path

        self._save_manifest()
        logger.info(f"Document {doc_id} extraction complete: {page_count} pages")

    def update_facts_count(self, doc_id: str, facts_count: int):
        """Update the count of facts extracted from this document."""
        with self._lock:
            if doc_id in self._documents:
                doc = self._documents[doc_id]
                doc.facts_extracted = facts_count
                doc.last_discovery_run = datetime.now().isoformat()

        self._save_manifest()

    def extract_text(self, doc_id: str) -> Optional[str]:
        """
        Extract text from a document.

        This is a hook that calls the appropriate extraction method
        based on file type. Stores result in extracted_text_path.

        Args:
            doc_id: Document ID

        Returns:
            Extracted text with [PAGE X] markers, or None if extraction fails
        """
        doc = self.get_document(doc_id)
        if not doc:
            logger.error(f"Document not found: {doc_id}")
            return None

        raw_path = Path(doc.raw_file_path)
        if not raw_path.exists():
            logger.error(f"Raw file not found: {raw_path}")
            self.update_status(doc_id, DocumentStatus.ERROR, "Raw file not found")
            return None

        # Update status
        self.update_status(doc_id, DocumentStatus.EXTRACTING)

        try:
            # Extract based on file type
            text = None
            page_count = 0

            if doc.mime_type == "application/pdf":
                text, page_count = self._extract_pdf(raw_path)
            elif doc.mime_type in (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ):
                text, page_count = self._extract_docx(raw_path)
            elif doc.mime_type in ("text/plain", "text/markdown", "text/csv"):
                text, page_count = self._extract_text_file(raw_path)
            elif doc.mime_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
                text, page_count = self._extract_xlsx(raw_path)
            else:
                # Try as text
                text, page_count = self._extract_text_file(raw_path)

            if text:
                # Save extracted text
                extracted_path = Path(doc.extracted_text_path)
                extracted_path.parent.mkdir(parents=True, exist_ok=True)
                with open(extracted_path, "w", encoding="utf-8") as f:
                    f.write(text)

                self.update_extraction_result(doc_id, page_count)
                logger.info(f"Extracted {page_count} pages from {doc.filename}")
                return text
            else:
                self.update_status(doc_id, DocumentStatus.ERROR, "No text extracted")
                return None

        except Exception as e:
            logger.error(f"Extraction failed for {doc_id}: {e}")
            self.update_status(doc_id, DocumentStatus.ERROR, str(e))
            return None

    def _extract_pdf(self, file_path: Path) -> tuple[str, int]:
        """Extract text from PDF with page markers."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.warning("PyMuPDF not installed, trying pdfplumber")
            try:
                import pdfplumber
                return self._extract_pdf_pdfplumber(file_path)
            except ImportError:
                raise ImportError("No PDF library available. Install pymupdf or pdfplumber.")

        doc = fitz.open(file_path)
        pages = []

        for page_num, page in enumerate(doc, 1):
            text = page.get_text()
            if text.strip():
                pages.append(f"[PAGE {page_num}]\n{text}")

        doc.close()
        return "\n\n".join(pages), len(pages)

    def _extract_pdf_pdfplumber(self, file_path: Path) -> tuple[str, int]:
        """Fallback PDF extraction using pdfplumber."""
        import pdfplumber

        pages = []
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                if text and text.strip():
                    pages.append(f"[PAGE {page_num}]\n{text}")

        return "\n\n".join(pages), len(pages)

    def _extract_docx(self, file_path: Path) -> tuple[str, int]:
        """Extract text from DOCX."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx not installed")

        doc = DocxDocument(file_path)
        paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # DOCX doesn't have clear page breaks, estimate
        text = "\n\n".join(paragraphs)
        # Rough estimate: ~3000 chars per page
        page_count = max(1, len(text) // 3000)

        return f"[PAGE 1]\n{text}", page_count

    def _extract_xlsx(self, file_path: Path) -> tuple[str, int]:
        """Extract text from Excel file."""
        try:
            import openpyxl
        except ImportError:
            raise ImportError("openpyxl not installed")

        wb = openpyxl.load_workbook(file_path, data_only=True)
        sheets = []

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            rows = []
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                if row_text.strip().replace("|", "").strip():
                    rows.append(row_text)

            if rows:
                sheets.append(f"[SHEET: {sheet_name}]\n" + "\n".join(rows))

        wb.close()
        return "\n\n".join(sheets), len(sheets)

    def _extract_text_file(self, file_path: Path) -> tuple[str, int]:
        """Extract text from plain text files."""
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()

        # Estimate pages
        page_count = max(1, len(text) // 3000)
        return f"[PAGE 1]\n{text}", page_count

    def get_extracted_text(self, doc_id: str) -> Optional[str]:
        """
        Get the extracted text for a document.

        Args:
            doc_id: Document ID

        Returns:
            Extracted text or None if not available
        """
        doc = self.get_document(doc_id)
        if not doc:
            return None

        extracted_path = Path(doc.extracted_text_path)
        if not extracted_path.exists():
            # Try to extract
            return self.extract_text(doc_id)

        with open(extracted_path, "r", encoding="utf-8") as f:
            return f.read()

    def _save_manifest(self):
        """Save document registry to manifest file."""
        with self._lock:
            manifest = {
                "version": "1.0",
                "updated_at": datetime.now().isoformat(),
                "document_count": len(self._documents),
                "documents": [doc.to_dict() for doc in self._documents.values()]
            }

        # Write atomically (write to temp, then rename)
        temp_path = self.manifest_path.with_suffix(".tmp")
        with open(temp_path, "w", encoding="utf-8") as f:
            json.dump(manifest, f, indent=2)

        temp_path.rename(self.manifest_path)
        logger.debug(f"Saved manifest with {len(self._documents)} documents")

    def _load_manifest(self):
        """Load document registry from manifest file."""
        if not self.manifest_path.exists():
            logger.debug("No manifest found, starting fresh")
            return

        try:
            with open(self.manifest_path, "r", encoding="utf-8") as f:
                manifest = json.load(f)

            for doc_data in manifest.get("documents", []):
                doc = Document.from_dict(doc_data)
                self._documents[doc.doc_id] = doc
                self._hash_index[doc.hash_sha256] = doc.doc_id

            logger.info(f"Loaded {len(self._documents)} documents from manifest")

        except Exception as e:
            logger.error(f"Error loading manifest: {e}")

    def get_statistics(self) -> Dict[str, Any]:
        """Get document store statistics."""
        with self._lock:
            docs = list(self._documents.values())

        return {
            "total_documents": len(docs),
            "by_entity": {
                "target": len([d for d in docs if d.entity == "target"]),
                "buyer": len([d for d in docs if d.entity == "buyer"])
            },
            "by_authority": {
                "data_room": len([d for d in docs if d.authority_level == 1]),
                "correspondence": len([d for d in docs if d.authority_level == 2]),
                "discussion_notes": len([d for d in docs if d.authority_level == 3])
            },
            "by_status": {
                status.value: len([d for d in docs if d.status == status.value])
                for status in DocumentStatus
            },
            "total_facts_extracted": sum(d.facts_extracted for d in docs),
            "total_size_bytes": sum(d.file_size_bytes for d in docs)
        }

    def clear(self):
        """Clear all documents (for testing)."""
        with self._lock:
            self._documents.clear()
            self._hash_index.clear()

        if self.manifest_path.exists():
            self.manifest_path.unlink()

        logger.info("Document store cleared")
