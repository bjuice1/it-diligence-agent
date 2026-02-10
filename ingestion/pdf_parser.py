"""
PDF Document Parser for IT Diligence

Extracts text from PDF documents while preserving structure
for effective analysis by domain agents.
"""
import fitz  # PyMuPDF
from pathlib import Path
from typing import Dict, List, Optional
from dataclasses import dataclass, field
import json
import re

from tools_v2.document_preprocessor import DocumentPreprocessor


@dataclass
class DocumentSection:
    """A section of extracted content"""
    title: str
    content: str
    page_numbers: List[int]
    section_type: str = "text"  # text, table, list


@dataclass
class ParsedDocument:
    """Structured representation of a parsed PDF"""
    filename: str
    total_pages: int
    sections: List[DocumentSection] = field(default_factory=list)
    raw_text: str = ""
    metadata: Dict = field(default_factory=dict)
    
    def to_dict(self) -> Dict:
        return {
            "filename": self.filename,
            "total_pages": self.total_pages,
            "sections": [
                {
                    "title": s.title,
                    "content": s.content,
                    "page_numbers": s.page_numbers,
                    "section_type": s.section_type
                }
                for s in self.sections
            ],
            "raw_text": self.raw_text,
            "metadata": self.metadata
        }
    
    def get_text_for_analysis(self) -> str:
        """Get formatted text optimized for LLM analysis"""
        output = []

        # Classify entity based on filename patterns
        entity = self._classify_entity_from_filename()
        if entity:
            output.append(f"# ENTITY: {entity.upper()}")
            output.append(f"# (This document describes the {entity} company)")

        output.append(f"# Document: {self.filename}")
        output.append(f"Pages: {self.total_pages}\n")

        if self.sections:
            for section in self.sections:
                if section.title:
                    output.append(f"\n## {section.title}")
                output.append(section.content)
        else:
            output.append(self.raw_text)

        return "\n".join(output)

    def _classify_entity_from_filename(self) -> Optional[str]:
        """Classify document as TARGET or BUYER based on filename patterns"""
        filename_lower = self.filename.lower()

        # Target patterns
        target_patterns = [
            'target', 'target_profile', 'target company',
            'acquisition target', 'seller', 'sellside'
        ]
        for pattern in target_patterns:
            if pattern in filename_lower:
                return 'target'

        # Buyer patterns
        buyer_patterns = [
            'buyer', 'buyer_profile', 'buyer company',
            'acquirer', 'parent', 'buyside'
        ]
        for pattern in buyer_patterns:
            if pattern in filename_lower:
                return 'buyer'

        return None


class PDFParser:
    """
    Parses PDF documents for IT diligence analysis.
    
    Optimized for:
    - IT infrastructure documentation
    - Network diagrams (extracts text annotations)
    - Security assessment reports
    - Vendor contracts and SLAs
    """
    
    # Patterns that indicate section headers in IT docs
    HEADER_PATTERNS = [
        r'^(?:SECTION\s+)?(\d+\.?\d*\.?\d*)\s+(.+)$',  # "1.2.3 Title" or "SECTION 1 Title"
        r'^([A-Z][A-Z\s&]+)$',  # "ALL CAPS HEADER"
        r'^(Executive Summary|Overview|Introduction|Infrastructure|Network|Security|Applications|Recommendations)',
        r'^(Current State|Future State|Gap Analysis|Risk Assessment|Cost Estimate)',
    ]
    
    # IT-specific terms to identify relevant sections
    IT_KEYWORDS = {
        'infrastructure': ['server', 'datacenter', 'data center', 'hosting', 'vmware', 'virtual', 
                          'storage', 'san', 'nas', 'cloud', 'aws', 'azure', 'gcp', 'compute'],
        'network': ['network', 'wan', 'lan', 'mpls', 'sd-wan', 'firewall', 'router', 'switch',
                   'vlan', 'dns', 'dhcp', 'vpn', 'bandwidth', 'latency', 'cisco'],
        'cybersecurity': ['security', 'cyber', 'identity', 'iam', 'mfa', 'sso', 'compliance',
                         'soc', 'pci', 'hipaa', 'gdpr', 'vulnerability', 'patch', 'encryption',
                         'firewall', 'siem', 'edr', 'antivirus', 'threat']
    }
    
    def __init__(self):
        self.documents: List[ParsedDocument] = []
        self._preprocessor = DocumentPreprocessor(aggressive=False)
    
    def parse_file(self, filepath: Path) -> ParsedDocument:
        """Parse a single PDF file"""
        filepath = Path(filepath)
        
        if not filepath.exists():
            raise FileNotFoundError(f"PDF not found: {filepath}")
        
        if filepath.suffix.lower() != '.pdf':
            raise ValueError(f"Not a PDF file: {filepath}")
        
        doc = fitz.open(filepath)
        
        parsed = ParsedDocument(
            filename=filepath.name,
            total_pages=len(doc),
            metadata={
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "created": doc.metadata.get("creationDate", ""),
            }
        )
        
        # Extract text page by page
        all_text = []

        for page_num, page in enumerate(doc, 1):
            page_text = page.get_text("text")
            page_text = self._preprocessor.clean(page_text)
            all_text.append(page_text)

        # Merge tables that span page boundaries before section detection
        merged_texts = self._merge_split_tables(all_text)

        # Build sections from merged page texts
        current_section = None
        current_content = []
        current_pages = []

        for page_num, page_text in enumerate(merged_texts, 1):
            # Try to identify sections
            lines = page_text.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue

                # Check if this line is a header
                is_header = self._is_header(line)

                if is_header and current_section:
                    # Save previous section
                    parsed.sections.append(DocumentSection(
                        title=current_section,
                        content='\n'.join(current_content),
                        page_numbers=current_pages
                    ))
                    current_section = line
                    current_content = []
                    current_pages = [page_num]
                elif is_header:
                    current_section = line
                    current_pages = [page_num]
                else:
                    current_content.append(line)
                    if page_num not in current_pages:
                        current_pages.append(page_num)

        # Don't forget the last section
        if current_section and current_content:
            parsed.sections.append(DocumentSection(
                title=current_section,
                content='\n'.join(current_content),
                page_numbers=current_pages
            ))

        parsed.raw_text = '\n\n'.join(merged_texts)
        
        doc.close()
        self.documents.append(parsed)
        
        return parsed
    
    def parse_directory(self, dirpath: Path) -> List[ParsedDocument]:
        """Parse all PDFs in a directory"""
        dirpath = Path(dirpath)
        results = []
        
        for pdf_file in dirpath.glob("*.pdf"):
            try:
                parsed = self.parse_file(pdf_file)
                results.append(parsed)
                print(f"✓ Parsed: {pdf_file.name} ({parsed.total_pages} pages)")
            except Exception as e:
                print(f"✗ Failed: {pdf_file.name} - {e}")
        
        return results
    
    def _merge_split_tables(self, pages_text: List[str]) -> List[str]:
        """Detect and merge tables that span page boundaries.

        Strategy: If a page ends with a partial markdown table row (line starts with '|'
        but the table has fewer columns than the previous table block), merge the
        continuation rows into the previous page's table block.

        Also handles: page ending mid-row (line ends with '|' and next page starts with '|')
        """
        if len(pages_text) < 2:
            return pages_text

        merged = [pages_text[0]]
        for i in range(1, len(pages_text)):
            prev = merged[-1]
            curr = pages_text[i]

            prev_lines = prev.rstrip().split('\n')
            curr_lines = curr.lstrip().split('\n')

            # Check if previous page ends inside a table
            prev_in_table = (
                len(prev_lines) >= 2
                and prev_lines[-1].strip().startswith('|')
            )

            # Check if current page starts with table continuation
            curr_continues_table = (
                len(curr_lines) >= 1
                and curr_lines[0].strip().startswith('|')
            )

            if prev_in_table and curr_continues_table:
                # Find where table continuation ends in current page
                table_end = 0
                for j, line in enumerate(curr_lines):
                    if line.strip().startswith('|'):
                        table_end = j + 1
                    else:
                        break

                # Merge table rows into previous page
                table_continuation = '\n'.join(curr_lines[:table_end])
                remaining = '\n'.join(curr_lines[table_end:])

                merged[-1] = prev + '\n' + table_continuation
                if remaining.strip():
                    merged.append(remaining)
            else:
                merged.append(curr)

        return merged

    def _is_header(self, line: str) -> bool:
        """Determine if a line is likely a section header"""
        # Too long to be a header
        if len(line) > 100:
            return False
        
        # Check against header patterns
        for pattern in self.HEADER_PATTERNS:
            if re.match(pattern, line, re.IGNORECASE):
                return True
        
        # Check if it's short and ends without punctuation (likely header)
        if len(line) < 60 and not line.endswith(('.', ',', ';', ':')):
            words = line.split()
            # Headers often have capitalized words
            if len(words) <= 8 and sum(1 for w in words if w[0].isupper()) >= len(words) * 0.5:
                return True
        
        return False
    
    def identify_domain_relevance(self, parsed: ParsedDocument) -> Dict[str, float]:
        """Score how relevant a document is to each domain"""
        text_lower = parsed.raw_text.lower()
        scores = {}
        
        for domain, keywords in self.IT_KEYWORDS.items():
            count = sum(text_lower.count(kw) for kw in keywords)
            # Normalize by document length
            scores[domain] = count / (len(text_lower) / 1000)  # per 1000 chars
        
        return scores
    
    def get_combined_text(self) -> str:
        """Get all parsed documents as a single text block for analysis"""
        output = []
        for doc in self.documents:
            output.append(doc.get_text_for_analysis())
            output.append("\n" + "="*80 + "\n")
        return '\n'.join(output)
    
    def save_extracted(self, output_path: Path):
        """Save extracted content to JSON for debugging/review"""
        output_path = Path(output_path)
        data = {
            "documents": [doc.to_dict() for doc in self.documents],
            "combined_text_length": len(self.get_combined_text())
        }
        with open(output_path, 'w') as f:
            json.dump(data, f, indent=2)
        print(f"Saved extracted content to: {output_path}")


def parse_documents(file_paths: List[Path]) -> List[Dict]:
    """
    Parse multiple documents and return list of dicts with content.

    Supports: PDF, TXT, MD, Markdown files

    Args:
        file_paths: List of file paths to parse

    Returns:
        List of dicts with 'filename' and 'content' keys
    """
    results = []
    parser = PDFParser()

    for filepath in file_paths:
        filepath = Path(filepath)

        if not filepath.exists():
            print(f"File not found: {filepath}")
            continue

        try:
            suffix = filepath.suffix.lower()

            if suffix == '.pdf':
                # Parse PDF
                parsed = parser.parse_file(filepath)
                results.append({
                    'filename': filepath.name,
                    'content': parsed.get_text_for_analysis(),
                    'pages': parsed.total_pages,
                    'type': 'pdf'
                })

            elif suffix in ['.txt', '.md', '.markdown']:
                # Read text files directly
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                results.append({
                    'filename': filepath.name,
                    'content': content,
                    'pages': 1,
                    'type': 'text'
                })

            elif suffix in ['.doc', '.docx']:
                # For Word docs, try to extract text if python-docx available
                try:
                    import docx
                    doc = docx.Document(filepath)
                    content = '\n'.join([para.text for para in doc.paragraphs])
                    results.append({
                        'filename': filepath.name,
                        'content': content,
                        'pages': 1,
                        'type': 'docx'
                    })
                except ImportError:
                    print(f"python-docx not installed, skipping: {filepath.name}")

            elif suffix in ['.xls', '.xlsx']:
                # For Excel, try basic extraction
                try:
                    import openpyxl
                    wb = openpyxl.load_workbook(filepath, data_only=True)
                    content_parts = []
                    for sheet_name in wb.sheetnames:
                        sheet = wb[sheet_name]
                        content_parts.append(f"## Sheet: {sheet_name}\n")
                        for row in sheet.iter_rows(values_only=True):
                            row_text = '\t'.join([str(c) if c else '' for c in row])
                            if row_text.strip():
                                content_parts.append(row_text)
                    content = '\n'.join(content_parts)
                    results.append({
                        'filename': filepath.name,
                        'content': content,
                        'pages': len(wb.sheetnames),
                        'type': 'xlsx'
                    })
                except ImportError:
                    print(f"openpyxl not available for: {filepath.name}")
                except Exception as e:
                    print(f"Error reading Excel file {filepath.name}: {e}")

            else:
                # Try to read as text anyway
                try:
                    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                    results.append({
                        'filename': filepath.name,
                        'content': content,
                        'pages': 1,
                        'type': 'unknown'
                    })
                except Exception as e:
                    print(f"Could not read {filepath.name}: {e}")

        except Exception as e:
            print(f"Error parsing {filepath.name}: {e}")

    return results


# Convenience function
def parse_pdfs(input_path: Path) -> str:
    """
    Parse PDF(s) and return combined text for analysis.
    
    Args:
        input_path: Path to a PDF file or directory of PDFs
    
    Returns:
        Combined text from all documents
    """
    parser = PDFParser()
    input_path = Path(input_path)
    
    if input_path.is_file():
        parser.parse_file(input_path)
    elif input_path.is_dir():
        parser.parse_directory(input_path)
    else:
        raise ValueError(f"Invalid path: {input_path}")
    
    return parser.get_combined_text()


if __name__ == "__main__":
    # Test with sample
    import sys
    if len(sys.argv) > 1:
        text = parse_pdfs(Path(sys.argv[1]))
        print(f"\nExtracted {len(text)} characters")
        print("\nFirst 2000 chars:")
        print(text[:2000])
