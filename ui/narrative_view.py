"""
Executive Narrative View

Streamlit UI component for displaying the executive narrative with:
- Collapsible/expandable sections
- Clickable citations linking to facts
- Export buttons (Markdown, Word)
- Proper table rendering
"""

import streamlit as st
import re
import json
from pathlib import Path
from typing import Dict, List, Optional
from datetime import datetime
import io


# =============================================================================
# CITATION RENDERING
# =============================================================================

def render_with_citations(text: str, fact_lookup: Optional[Dict] = None) -> str:
    """
    Convert citations like (F-CYBER-003) to formatted references.

    If fact_lookup is provided, adds hover tooltips with fact content.
    """
    if not text:
        return ""

    # Pattern for fact citations: (F-DOMAIN-###)
    citation_pattern = r'\(F-([A-Z]+)-(\d{3})\)'

    def replace_citation(match):
        domain = match.group(1)
        num = match.group(2)
        fact_id = f"F-{domain}-{num}"

        if fact_lookup and fact_id in fact_lookup:
            fact = fact_lookup[fact_id]
            # Return with tooltip
            return f'<span class="citation" title="{fact.get("description", "")[:100]}...">[{fact_id}]</span>'
        else:
            return f'<span class="citation">[{fact_id}]</span>'

    return re.sub(citation_pattern, replace_citation, text)


def render_inference_labels(text: str) -> str:
    """
    Highlight inference labels in the text.
    """
    if not text:
        return ""

    # Highlight "Inference:" prefix
    text = re.sub(
        r'(Inference:)',
        r'<span class="inference-label">⚡ \1</span>',
        text,
        flags=re.IGNORECASE
    )

    # Highlight "Pattern:" prefix
    text = re.sub(
        r'(Pattern:)',
        r'<span class="pattern-label">🔄 \1</span>',
        text,
        flags=re.IGNORECASE
    )

    # Highlight "(GAP)" flags
    text = re.sub(
        r'\(GAP\)',
        r'<span class="gap-flag">⚠️ GAP</span>',
        text,
        flags=re.IGNORECASE
    )

    return text


# =============================================================================
# SECTION RENDERERS
# =============================================================================

def render_executive_summary(summary: List[str], expanded: bool = True):
    """Render the executive summary section."""
    with st.expander("📋 1) Executive Summary", expanded=expanded):
        if summary:
            for bullet in summary:
                formatted = render_inference_labels(bullet)
                st.markdown(f"• {formatted}", unsafe_allow_html=True)
        else:
            st.warning("Executive summary not available")

        st.caption(f"📊 {len(summary)} key points")


def render_org_structure(narrative: str, expanded: bool = False):
    """Render the organization structure narrative section."""
    with st.expander("🏢 2) How the IT Organization is Built", expanded=expanded):
        if narrative:
            # Split into paragraphs and render
            paragraphs = narrative.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    formatted = render_inference_labels(para)
                    st.markdown(formatted, unsafe_allow_html=True)
                    st.write("")  # Add spacing
        else:
            st.warning("Organization structure narrative not available")


def render_team_stories(stories: List[Dict], expanded: bool = False):
    """Render the team-by-team stories section."""
    with st.expander(f"👥 3) Team-by-Team Story ({len(stories)} teams)", expanded=expanded):
        if stories:
            for story in stories:
                func_name = story.get('function_name', 'Unknown Team')

                st.markdown(f"### {func_name}")

                # Day-to-day
                day_to_day = story.get('day_to_day', '')
                if day_to_day:
                    st.markdown(f"**What they do day-to-day:** {day_to_day}")

                # Strengths and Constraints in columns
                col1, col2 = st.columns(2)

                with col1:
                    st.markdown("**What they likely do well:**")
                    strengths = story.get('strengths', [])
                    if strengths:
                        for s in strengths:
                            st.markdown(f"• ✅ {s}")
                    else:
                        st.caption("No strengths identified")

                with col2:
                    st.markdown("**Where they're likely constrained:**")
                    constraints = story.get('constraints', [])
                    if constraints:
                        for c in constraints:
                            st.markdown(f"• ⚠️ {c}")
                    else:
                        st.caption("No constraints identified")

                # Dependencies
                upstream = story.get('upstream_dependencies', [])
                downstream = story.get('downstream_dependents', [])

                if upstream or downstream:
                    st.markdown("**Key dependencies:**")
                    if upstream:
                        st.markdown(f"• *Upstream:* {', '.join(upstream)}")
                    if downstream:
                        st.markdown(f"• *Downstream:* {', '.join(downstream)}")

                # M&A Implication
                mna_impl = story.get('mna_implication', '')
                if mna_impl:
                    st.info(f"**M&A Implication:** {mna_impl}")

                st.markdown("---")
        else:
            st.warning("No team stories available")


def render_mna_lens(mna_section: Dict, deal_type: str = "acquisition", expanded: bool = False):
    """Render the M&A lens section."""
    with st.expander("🔍 4) M&A Lens: Day-1 + TSA + Separation", expanded=expanded):
        if not mna_section:
            st.warning("M&A lens section not available")
            return

        # TSA-Exposed Functions
        tsa_funcs = mna_section.get('tsa_exposed_functions', [])
        if tsa_funcs:
            st.markdown("### TSA-Exposed Functions")
            for func in tsa_funcs:
                st.markdown(f"• {func}")
            st.write("")

        # Day-1 Requirements
        day_1 = mna_section.get('day_1_requirements', [])
        if day_1:
            st.markdown("### Day-1 Requirements")
            for req in day_1:
                st.markdown(f"• 🚨 {req}")
            st.write("")

        # Separation Considerations
        sep_considerations = mna_section.get('separation_considerations', [])
        if sep_considerations:
            st.markdown("### Separation Considerations")
            for i, consideration in enumerate(sep_considerations, 1):
                st.markdown(f"{i}. {consideration}")

        # Deal type context
        st.caption(f"📌 Deal Type: {deal_type.upper()}")


def render_benchmarks(statements: List[str], expanded: bool = False):
    """Render the benchmarks section."""
    with st.expander(f"📈 5) Benchmarks & Operating Posture ({len(statements)} statements)", expanded=expanded):
        if statements:
            for statement in statements:
                # Check if it's an inference
                if statement.lower().startswith('inference:'):
                    st.info(f"⚡ {statement}")
                else:
                    st.markdown(f"• {statement}")
        else:
            st.warning("No benchmark statements available")


def render_risks_table(risks: List[Dict], expanded: bool = False):
    """Render the risks table."""
    with st.expander(f"⚠️ 6a) Risks ({len(risks)} identified)", expanded=expanded):
        if risks:
            # Create table header
            st.markdown("""
            | Risk | Why it matters | Likely impact | Mitigation |
            |------|----------------|---------------|------------|
            """)

            # Add rows
            for risk in risks:
                r = risk.get('risk', '')
                why = risk.get('why_it_matters', '')
                impact = risk.get('likely_impact', '')
                mitigation = risk.get('mitigation', '')
                st.markdown(f"| {r} | {why} | {impact} | {mitigation} |")

            # Also show as cards for better readability
            st.markdown("---")
            st.markdown("**Detailed View:**")

            for risk in risks:
                severity = risk.get('severity', 'medium')
                severity_icon = {'critical': '🔴', 'high': '🟠', 'medium': '🟡', 'low': '🟢'}.get(severity, '⚪')

                with st.container():
                    st.markdown(f"{severity_icon} **{risk.get('risk', 'Unknown Risk')}**")
                    col1, col2 = st.columns(2)
                    with col1:
                        st.caption(f"Why it matters: {risk.get('why_it_matters', 'N/A')}")
                        st.caption(f"Likely impact: {risk.get('likely_impact', 'N/A')}")
                    with col2:
                        st.caption(f"Mitigation: {risk.get('mitigation', 'N/A')}")
                    st.markdown("---")
        else:
            st.warning("No risks identified")


def render_synergies_table(synergies: List[Dict], deal_type: str = "acquisition", expanded: bool = False):
    """Render the synergies/opportunities table."""
    title = "Synergies" if deal_type.lower() == "acquisition" else "Opportunities"

    with st.expander(f"💡 6b) {title} ({len(synergies)} identified)", expanded=expanded):
        if synergies:
            # Create table header
            st.markdown("""
            | Opportunity | Why it matters | Value mechanism | First step |
            |-------------|----------------|-----------------|------------|
            """)

            # Add rows
            for syn in synergies:
                opp = syn.get('opportunity', '')
                why = syn.get('why_it_matters', '')
                value = syn.get('value_mechanism', '')
                first = syn.get('first_step', '')
                st.markdown(f"| {opp} | {why} | {value} | {first} |")

            # Detailed cards
            st.markdown("---")
            st.markdown("**Detailed View:**")

            for syn in synergies:
                mechanism_type = syn.get('mechanism_type', 'cost_elimination')
                mechanism_icon = {
                    'cost_elimination': '🗑️',
                    'cost_avoidance': '🛡️',
                    'efficiency_gain': '⚡',
                    'capability_gain': '🎯',
                    'revenue_enablement': '📈'
                }.get(mechanism_type, '💡')

                with st.container():
                    st.markdown(f"{mechanism_icon} **{syn.get('opportunity', 'Unknown')}**")
                    st.caption(f"Value mechanism: {syn.get('value_mechanism', 'N/A')}")
                    st.caption(f"First step: {syn.get('first_step', 'N/A')}")
                    if syn.get('estimated_value'):
                        st.caption(f"Estimated value: {syn.get('estimated_value')}")
                    st.markdown("---")
        else:
            st.info(f"No {title.lower()} identified")


# =============================================================================
# EXPORT FUNCTIONS
# =============================================================================

def export_to_markdown(narrative: Dict) -> str:
    """Export narrative to markdown format."""
    md = f"# IT Operating Model Narrative ({narrative.get('company_name', 'Target Company')})\n\n"
    md += f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}*\n"
    md += f"*Deal Type: {narrative.get('deal_type', 'acquisition')}*\n\n"

    # Executive Summary
    md += "## 1) Executive Summary\n\n"
    for bullet in narrative.get('executive_summary', []):
        md += f"- {bullet}\n"
    md += "\n"

    # Org Structure
    md += "## 2) How the IT Organization is Built\n\n"
    md += narrative.get('org_structure_narrative', '') + "\n\n"

    # Team Stories
    md += "## 3) Team-by-Team Story\n\n"
    for story in narrative.get('team_stories', []):
        md += f"### {story.get('function_name', 'Unknown')}\n\n"
        md += f"**What they do day-to-day**: {story.get('day_to_day', '')}\n\n"
        md += "**What they likely do well**:\n"
        for s in story.get('strengths', []):
            md += f"- {s}\n"
        md += "\n**Where they're likely constrained**:\n"
        for c in story.get('constraints', []):
            md += f"- {c}\n"
        md += "\n**Key dependencies**:\n"
        md += f"- *Upstream*: {', '.join(story.get('upstream_dependencies', []))}\n"
        md += f"- *Downstream*: {', '.join(story.get('downstream_dependents', []))}\n"
        md += f"\n**M&A Implication**: {story.get('mna_implication', '')}\n\n"

    # M&A Lens
    md += "## 4) M&A Lens: Day-1 + TSA + Separation\n\n"
    mna = narrative.get('mna_lens_section', {})

    if mna.get('tsa_exposed_functions'):
        md += "### TSA-Exposed Functions\n"
        for func in mna['tsa_exposed_functions']:
            md += f"- {func}\n"
        md += "\n"

    if mna.get('day_1_requirements'):
        md += "### Day-1 Requirements\n"
        for req in mna['day_1_requirements']:
            md += f"- {req}\n"
        md += "\n"

    if mna.get('separation_considerations'):
        md += "### Separation Considerations\n"
        for i, c in enumerate(mna['separation_considerations'], 1):
            md += f"{i}. {c}\n"
        md += "\n"

    # Benchmarks
    md += "## 5) Benchmarks & Operating Posture\n\n"
    for statement in narrative.get('benchmark_statements', []):
        md += f"- {statement}\n"
    md += "\n"

    # Risks Table
    md += "## 6) Risks and Synergies\n\n"
    md += "### Risks\n\n"
    md += "| Risk | Why it matters | Likely impact | Mitigation |\n"
    md += "|------|----------------|---------------|------------|\n"
    for risk in narrative.get('risks_table', []):
        md += f"| {risk.get('risk', '')} | {risk.get('why_it_matters', '')} | {risk.get('likely_impact', '')} | {risk.get('mitigation', '')} |\n"
    md += "\n"

    # Synergies Table
    md += "### Synergies / Opportunities\n\n"
    md += "| Opportunity | Why it matters | Value mechanism | First step |\n"
    md += "|-------------|----------------|-----------------|------------|\n"
    for syn in narrative.get('synergies_table', []):
        md += f"| {syn.get('opportunity', '')} | {syn.get('why_it_matters', '')} | {syn.get('value_mechanism', '')} | {syn.get('first_step', '')} |\n"

    return md


def export_to_word(narrative: Dict) -> Optional[bytes]:
    """
    Export narrative to Word format.
    Requires python-docx library.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return None

    doc = Document()

    # Title
    title = doc.add_heading("IT Operating Model Narrative", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(f"{narrative.get('company_name', 'Target Company')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph(f"Deal Type: {narrative.get('deal_type', 'acquisition').upper()}")

    # Executive Summary
    doc.add_heading("1) Executive Summary", level=1)
    for bullet in narrative.get('executive_summary', []):
        doc.add_paragraph(bullet, style='List Bullet')

    # Org Structure
    doc.add_heading("2) How the IT Organization is Built", level=1)
    doc.add_paragraph(narrative.get('org_structure_narrative', ''))

    # Team Stories
    doc.add_heading("3) Team-by-Team Story", level=1)
    for story in narrative.get('team_stories', []):
        doc.add_heading(story.get('function_name', 'Unknown'), level=2)
        doc.add_paragraph(f"What they do day-to-day: {story.get('day_to_day', '')}")

        doc.add_paragraph("What they likely do well:")
        for s in story.get('strengths', []):
            doc.add_paragraph(s, style='List Bullet')

        doc.add_paragraph("Where they're likely constrained:")
        for c in story.get('constraints', []):
            doc.add_paragraph(c, style='List Bullet')

        doc.add_paragraph(f"M&A Implication: {story.get('mna_implication', '')}")

    # M&A Lens
    doc.add_heading("4) M&A Lens: Day-1 + TSA + Separation", level=1)
    mna = narrative.get('mna_lens_section', {})

    if mna.get('separation_considerations'):
        doc.add_heading("Separation Considerations", level=2)
        for c in mna['separation_considerations']:
            doc.add_paragraph(c, style='List Number')

    # Benchmarks
    doc.add_heading("5) Benchmarks & Operating Posture", level=1)
    for statement in narrative.get('benchmark_statements', []):
        doc.add_paragraph(statement, style='List Bullet')

    # Risks
    doc.add_heading("6) Risks and Synergies", level=1)
    doc.add_heading("Risks", level=2)

    # Create risks table
    risks = narrative.get('risks_table', [])
    if risks:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        header = table.rows[0].cells
        header[0].text = "Risk"
        header[1].text = "Why it matters"
        header[2].text = "Likely impact"
        header[3].text = "Mitigation"

        for risk in risks:
            row = table.add_row().cells
            row[0].text = risk.get('risk', '')
            row[1].text = risk.get('why_it_matters', '')
            row[2].text = risk.get('likely_impact', '')
            row[3].text = risk.get('mitigation', '')

    # Synergies
    doc.add_heading("Synergies / Opportunities", level=2)
    synergies = narrative.get('synergies_table', [])
    if synergies:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        header = table.rows[0].cells
        header[0].text = "Opportunity"
        header[1].text = "Why it matters"
        header[2].text = "Value mechanism"
        header[3].text = "First step"

        for syn in synergies:
            row = table.add_row().cells
            row[0].text = syn.get('opportunity', '')
            row[1].text = syn.get('why_it_matters', '')
            row[2].text = syn.get('value_mechanism', '')
            row[3].text = syn.get('first_step', '')

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# =============================================================================
# MAIN RENDER FUNCTION
# =============================================================================

def render_narrative_section(
    narrative: Dict,
    deal_type: str = "acquisition",
    fact_lookup: Optional[Dict] = None,
    show_exports: bool = True
):
    """
    Main function to render the full executive narrative section.

    Args:
        narrative: The narrative dict (from ExecutiveNarrative.to_dict())
        deal_type: Type of deal for context
        fact_lookup: Optional dict mapping fact IDs to fact content
        show_exports: Whether to show export buttons
    """
    company_name = narrative.get('company_name', 'Target Company')

    # Header
    st.markdown(f"## 📄 Executive Narrative: {company_name}")
    st.caption(f"Deal Type: {deal_type.upper()} | Generated: {narrative.get('created_at', 'N/A')}")

    # Inject custom CSS for citations
    st.markdown("""
    <style>
    .citation {
        background-color: #e8f4f8;
        padding: 2px 6px;
        border-radius: 4px;
        font-family: monospace;
        font-size: 0.85em;
        color: #0066cc;
        cursor: pointer;
    }
    .inference-label {
        background-color: #fff3cd;
        padding: 2px 6px;
        border-radius: 4px;
        font-weight: 600;
    }
    .pattern-label {
        background-color: #d1ecf1;
        padding: 2px 6px;
        border-radius: 4px;
        font-weight: 600;
    }
    .gap-flag {
        background-color: #f8d7da;
        padding: 2px 6px;
        border-radius: 4px;
        font-weight: 600;
        color: #721c24;
    }
    </style>
    """, unsafe_allow_html=True)

    # Quality metrics (if available)
    if 'validation' in narrative:
        validation = narrative['validation']
        score = validation.get('score', 0)
        status = "✅ PASS" if validation.get('valid', False) else "⚠️ NEEDS REVIEW"

        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Quality Score", f"{score}/100")
        with col2:
            st.metric("Status", status)
        with col3:
            issues = len(validation.get('issues', []))
            st.metric("Issues", issues)

        st.markdown("---")

    # Render all sections
    render_executive_summary(narrative.get('executive_summary', []), expanded=True)
    render_org_structure(narrative.get('org_structure_narrative', ''))
    render_team_stories(narrative.get('team_stories', []))
    render_mna_lens(narrative.get('mna_lens_section', {}), deal_type)
    render_benchmarks(narrative.get('benchmark_statements', []))
    render_risks_table(narrative.get('risks_table', []))
    render_synergies_table(narrative.get('synergies_table', []), deal_type)

    # Export buttons
    if show_exports:
        st.markdown("---")
        st.markdown("### 📥 Export Options")

        col1, col2, col3 = st.columns(3)

        with col1:
            # Markdown export
            md_content = export_to_markdown(narrative)
            st.download_button(
                label="📝 Download Markdown",
                data=md_content,
                file_name=f"narrative_{company_name.lower().replace(' ', '_')}.md",
                mime="text/markdown"
            )

        with col2:
            # Word export
            docx_content = export_to_word(narrative)
            if docx_content:
                st.download_button(
                    label="📄 Download Word",
                    data=docx_content,
                    file_name=f"narrative_{company_name.lower().replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.button("📄 Word Export", disabled=True, help="Install python-docx for Word export")

        with col3:
            # JSON export
            json_content = json.dumps(narrative, indent=2)
            st.download_button(
                label="🔧 Download JSON",
                data=json_content,
                file_name=f"narrative_{company_name.lower().replace(' ', '_')}.json",
                mime="application/json"
            )


def load_narrative_from_file(file_path: Path) -> Optional[Dict]:
    """Load narrative from a JSON file."""
    try:
        with open(file_path, 'r') as f:
            return json.load(f)
    except Exception as e:
        st.error(f"Failed to load narrative: {e}")
        return None


# =============================================================================
# EXPORTS
# =============================================================================

__all__ = [
    'render_narrative_section',
    'render_with_citations',
    'export_to_markdown',
    'export_to_word',
    'load_narrative_from_file'
]
